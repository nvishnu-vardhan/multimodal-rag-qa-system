import streamlit as st
import os
import tempfile
from typing import List
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
from io import BytesIO

# OpenAI and LangChain
from openai import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.schema import Document as LangchainDocument

# Web Search
from duckduckgo_search import DDGS

# Configure page
st.set_page_config(
    page_title="Multi-Modal RAG QA System",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
    .main-header {font-size: 2.5rem; font-weight: 700; color: #1f77b4; text-align: center; padding: 1rem 0;}
    .sub-header {font-size: 1.2rem; color: #555; text-align: center; margin-bottom: 2rem;}
    .file-item {background: #f0f2f6; padding: 10px; margin: 5px 0; border-radius: 5px; border-left: 4px solid #1f77b4;}
    .citation-box {background-color: #f0f2f6; border-left: 4px solid #1f77b4; padding: 10px; margin: 10px 0; border-radius: 5px;}
    </style>
""", unsafe_allow_html=True)

st.markdown('<h1 class="main-header">🤖 Multi-Modal RAG QA System</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Advanced Document Intelligence with Text, Tables, Images & Citations</p>', unsafe_allow_html=True)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'documents_processed' not in st.session_state:
    st.session_state.documents_processed = False
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'uploaded_file_names' not in st.session_state:
    st.session_state.uploaded_file_names = []

# Document processing functions
def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting PDF: {str(e)}")
        return ""

def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(BytesIO(file.read()))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error extracting DOCX: {str(e)}")
        return ""

def extract_text_from_image(file) -> str:
    """Extract text from image using OCR"""
    try:
        image = Image.open(BytesIO(file.read()))
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"Error extracting from image: {str(e)}")
        return ""

def process_documents(files, api_key) -> Chroma:
    """Process uploaded documents and create vector store"""
    all_texts = []
    
    for file in files:
        file_name = file.name
        file_type = file.type
        
        st.info(f"Processing: {file_name}")
        
        text = ""
        if file_type == "application/pdf":
            text = extract_text_from_pdf(file)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            text = extract_text_from_docx(file)
        elif file_type in ["image/png", "image/jpeg", "image/jpg"]:
            text = extract_text_from_image(file)
        
        if text:
            all_texts.append({"text": text, "source": file_name})
    
    if not all_texts:
        st.error("No text extracted from documents")
        return None
    
    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    
    documents = []
    for doc_data in all_texts:
        chunks = text_splitter.split_text(doc_data["text"])
        for chunk in chunks:
            documents.append(
                LangchainDocument(
                    page_content=chunk,
                    metadata={"source": doc_data["source"]}
                )
            )
    
    # Create embeddings and vector store
    embeddings = OpenAIEmbeddings(openai_api_key=api_key)
    vector_store = Chroma.from_documents(documents, embeddings)
    
    return vector_store

def search_web(query: str, num_results: int = 3) -> List[str]:
    """Search the web using DuckDuckGo"""
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=num_results))
            return [f"[{r['title']}]({r['link']}): {r['body']}" for r in results]
    except Exception as e:
        st.error(f"Web search error: {str(e)}")
        return []

def generate_response(query: str, vector_store, api_key: str, use_web_search: bool = False) -> tuple:
    """Generate response using RAG"""
    client = OpenAI(api_key=api_key)
    
    # Retrieve relevant documents
    docs = vector_store.similarity_search(query, k=3)
    context = "\n\n".join([f"From {doc.metadata['source']}:\n{doc.page_content}" for doc in docs])
    
    # Web search if enabled
    web_context = ""
    if use_web_search:
        web_results = search_web(query)
        if web_results:
            web_context = "\n\nWeb Search Results:\n" + "\n".join(web_results)
    
    # Generate response
    prompt = f"""Based on the following information, answer the user's question accurately and concisely.

Document Context:
{context}
{web_context}

User Question: {query}

Provide a clear answer with citations from the source documents."""
    
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    
    answer = response.choices[0].message.content
    sources = [doc.metadata['source'] for doc in docs]
    
    return answer, sources

# Sidebar
with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("OpenAI API Key", type="password", help="Enter your OpenAI API key")
    
    st.divider()
    
    st.header("📄 Document Upload")
    uploaded_files = st.file_uploader(
        "Upload documents (PDF, DOCX, Images)",
        type=['pdf', 'docx', 'png', 'jpg', 'jpeg'],
        accept_multiple_files=True
    )
    
    # Show uploaded files
    if uploaded_files:
        st.write(f"**{len(uploaded_files)} file(s) selected:**")
        for file in uploaded_files:
            st.markdown(f'<div class="file-item">📄 {file.name} ({file.size/1024:.1f}KB)</div>', unsafe_allow_html=True)
    
    if st.button("🔄 Process Documents", use_container_width=True):
        if not uploaded_files:
            st.warning("Please upload documents first")
        elif not api_key:
            st.warning("Please enter your OpenAI API key")
        else:
            with st.spinner("Processing documents..."):
                vector_store = process_documents(uploaded_files, api_key)
                if vector_store:
                    st.session_state.vector_store = vector_store
                    st.session_state.documents_processed = True
                    st.session_state.uploaded_file_names = [f.name for f in uploaded_files]
                    st.success(f"✅ Processed {len(uploaded_files)} document(s)!")
    
    st.divider()
    
    # Web Search Toggle
    st.header("🔍 Search Options")
    use_web_search = st.checkbox("Enable Web Search", value=False, help="Search the web if document context is insufficient")
    
    st.divider()
    
    if st.session_state.documents_processed:
        st.header("📊 Statistics")
        st.metric("Documents", len(st.session_state.uploaded_file_names))
        st.metric("Queries", len(st.session_state.chat_history) // 2)

# Main content
main_tab, docs_tab, about_tab = st.tabs(["💬 Chat", "📚 Documents", "ℹ️ About"])

with main_tab:
    # Display chat history
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            if "sources" in message:
                with st.expander("📖 View Sources"):
                    for source in message["sources"]:
                        st.markdown(f'<div class="citation-box">📄 {source}</div>', unsafe_allow_html=True)

with docs_tab:
    st.header("📚 Processed Documents")
    
    if st.session_state.documents_processed and st.session_state.uploaded_file_names:
        st.success(f"✅ {len(st.session_state.uploaded_file_names)} document(s) processed successfully!")
        for idx, filename in enumerate(st.session_state.uploaded_file_names, 1):
            st.markdown(f'<div class="file-item">{idx}. {filename}</div>', unsafe_allow_html=True)
    else:
        st.info("No documents uploaded yet. Upload documents using the sidebar.")

with about_tab:
    st.header("ℹ️ About Multi-Modal RAG QA System")
    st.markdown("""
    ### 🎯 Features
    
    **Fully Functional RAG System with:**
    - ✅ PDF text extraction
    - ✅ DOCX document processing
    - ✅ Image OCR (Optical Character Recognition)
    - ✅ Vector database (ChromaDB)
    - ✅ OpenAI-powered question answering
    - ✅ Web search integration (DuckDuckGo)
    - ✅ Source attribution and citations
    
    ### 🚀 How to Use
    
    1. Enter your OpenAI API key in the sidebar
    2. Upload your documents (PDF, DOCX, or Images)
    3. Click "Process Documents" to extract and index content
    4. Ask questions in the Chat tab
    5. Enable "Web Search" for additional context from the internet
    6. View source citations for each answer
    
    ### 🔗 Repository
    
    [GitHub Repository](https://github.com/nvishnu-vardhan/multimodal-rag-qa-system)
    
    **Built with ❤️ for advanced document intelligence**
    """)

# Chat input (outside tabs for Streamlit API compliance)
if prompt := st.chat_input("Ask a question about your documents..."):
    if not st.session_state.documents_processed:
        st.warning("Please upload and process documents first!")
    elif not api_key:
        st.warning("Please enter your OpenAI API key in the sidebar!")
    else:
        # Add user message
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                answer, sources = generate_response(
                    prompt, 
                    st.session_state.vector_store, 
                    api_key,
                    use_web_search
                )
                
                st.markdown(answer)
                
                with st.expander("📖 View Sources"):
                    for source in sources:
                        st.markdown(f'<div class="citation-box">📄 {source}</div>', unsafe_allow_html=True)
                
                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": answer,
                    "sources": sources
                })
